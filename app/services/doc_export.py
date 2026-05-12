"""Binary document exporters for a conclave task: PDF and DOCX.

Pure rendering: each function takes the same structured task data that the
markdown exporter takes (task envelope + messages + final_result + agent_runs)
and returns bytes. No DB, no filesystem I/O - the caller streams the bytes.

The markdown / plain-text path lives in `exporter.export_to_markdown`; this
module is for the formats that need a binary container. Layout mirrors the
markdown exporter's section order so all four formats are recognizable as the
same document: header -> question -> decision -> final result -> transcript.
"""

from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from typing import Any

# Both libraries are hard deps of this service (see requirements.txt). Import at
# module load so a missing install fails loudly at startup, not mid-request.
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import html as _html


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_COPYRIGHT_LINE = "Generated from AI Switchboard. Copyright © 2026 digitalgods.ai. All rights reserved."


def filename_stem(task: dict) -> str:
    """A filesystem-safe filename stem like 'conclave-is-there-a-god-tsk_01KR...'.

    The caller appends the format extension. Browsers' Save dialog uses this as
    the suggested filename, so keep it readable but conservative.
    """
    mode = (task.get("mode") or "task").lower()
    q = (task.get("user_request") or "").strip().lower()
    slug = re.sub(r"[^a-z0-9]+", "-", q).strip("-")[:48].strip("-")
    tid = task.get("id") or "unknown"
    stem = f"{mode}-{slug}-{tid}" if slug else f"{mode}-{tid}"
    # Final guard: nothing exotic.
    return re.sub(r"[^A-Za-z0-9._-]+", "-", stem)


def _agents_summary(task: dict) -> str:
    parts: list[str] = []
    if task.get("primary_agent"):
        parts.append(f"{task['primary_agent']} (primary)")
    for c in (task.get("consultants") or []):
        parts.append(str(c))
    return ", ".join(parts)


def _structured_fields(m: dict) -> list[tuple[str, str]]:
    """Return (label, value) pairs from a message's structured payload, in a
    stable, human-friendly order. Skips empty values and bookkeeping keys."""
    s = m.get("structured")
    out: list[tuple[str, str]] = []
    if not isinstance(s, dict):
        content = m.get("content")
        if isinstance(content, str) and content.strip():
            out.append(("content", content))
        return out
    skip = {"protocol_version", "task_id", "agent", "role", "message_type"}
    # Preferred order; anything else trails alphabetically.
    order = ["convergence", "agreement", "resolution_status", "summary",
             "position", "critique", "analysis", "user_input_question"]
    seen = set()
    for key in order:
        v = s.get(key)
        if _nonempty(v):
            out.append((key, _stringify(v)))
            seen.add(key)
    for key in sorted(s.keys()):
        if key in skip or key in seen:
            continue
        v = s.get(key)
        if _nonempty(v):
            out.append((key, _stringify(v)))
    return out


def _nonempty(v: Any) -> bool:
    if v is None:
        return False
    if isinstance(v, str):
        return v.strip() != ""
    if isinstance(v, (list, dict, tuple, set)):
        return len(v) > 0
    return True


def _stringify(v: Any) -> str:
    if isinstance(v, str):
        return v
    if isinstance(v, (int, float, bool)):
        return str(v)
    if isinstance(v, list):
        return "\n".join(f"- {_stringify(x)}" for x in v)
    if isinstance(v, dict):
        return "\n".join(f"{k}: {_stringify(val)}" for k, val in v.items())
    return str(v)


def _prettify(key: str) -> str:
    return key.replace("_", " ").strip().upper()


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------

def render_pdf(task: dict, messages: list[dict], final_result: dict | None,
               agent_runs: list[dict]) -> bytes:
    def esc(s: Any) -> str:
        if s is None:
            return ""
        return _html.escape(str(s)).replace("\n", "<br/>")

    styles = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=styles["Title"], fontSize=20, leading=24, spaceAfter=6)
    SUB = ParagraphStyle("SUB", parent=styles["Normal"], fontSize=10, leading=13,
                         textColor=HexColor("#555555"), spaceAfter=2)
    META = ParagraphStyle("META", parent=styles["Normal"], fontSize=9, leading=12,
                          textColor=HexColor("#666666"))
    SECTION = ParagraphStyle("SECTION", parent=styles["Heading2"], fontSize=14, leading=18,
                             spaceBefore=14, spaceAfter=6, textColor=HexColor("#1a1a1a"))
    TURN = ParagraphStyle("TURN", parent=styles["Heading3"], fontSize=11, leading=14,
                          spaceBefore=10, spaceAfter=2, textColor=HexColor("#2c5777"))
    LABEL = ParagraphStyle("LABEL", parent=styles["Normal"], fontSize=8, leading=11,
                           textColor=HexColor("#888888"), spaceBefore=4)
    BODY = ParagraphStyle("BODY", parent=styles["Normal"], fontSize=10, leading=14,
                          alignment=TA_LEFT, spaceAfter=2)
    FOOTER = ParagraphStyle("FOOTER", parent=styles["Normal"], fontSize=8, leading=10,
                            textColor=HexColor("#999999"))

    story: list = []
    story.append(Paragraph("AI Conclave — Deliberation Record", H1))
    story.append(Paragraph(esc(task.get("user_request")), SUB))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=0.7, color=HexColor("#cccccc")))
    story.append(Spacer(1, 6))

    for line in [
        f"<b>Task ID:</b> {esc(task.get('id'))}",
        f"<b>Mode:</b> {esc(task.get('mode'))} &nbsp;&nbsp; <b>Status:</b> {esc(task.get('status'))}",
        f"<b>Participants:</b> {esc(_agents_summary(task))}",
        f"<b>Invoked by:</b> {esc(task.get('source_agent'))}",
        f"<b>Created:</b> {esc(task.get('created_at'))}",
        f"<b>Updated:</b> {esc(task.get('updated_at'))}",
    ]:
        if "None" in line and ("Invoked by" in line):
            continue
        story.append(Paragraph(line, META))
    story.append(Spacer(1, 4))
    story.append(HRFlowable(width="100%", thickness=0.7, color=HexColor("#cccccc")))

    # Decision (if any)
    decision = task.get("user_decision")
    if isinstance(decision, str) and decision.strip():
        story.append(Paragraph("Recorded Decision", SECTION))
        if task.get("user_decided_at"):
            story.append(Paragraph(f"<i>Recorded at {esc(task['user_decided_at'])}</i>", META))
        story.append(Paragraph(esc(decision), BODY))

    # Final result
    story.append(Paragraph("Final Result", SECTION))
    fr = final_result or {}
    if fr:
        story.append(Paragraph(f"<b>Agreement level:</b> {esc(fr.get('agreement_level'))}", BODY))
        if fr.get("resolution_status"):
            story.append(Paragraph(f"<b>Resolution status:</b> {esc(fr.get('resolution_status'))}", BODY))
        story.append(Spacer(1, 4))
        story.append(Paragraph(esc(fr.get("final_answer")), BODY))
        for dd in (fr.get("disagreements") or []):
            story.append(Paragraph("Disagreements", TURN))
            story.append(Paragraph(f"<b>{esc(dd.get('topic'))}</b>", BODY))
            if dd.get("primary_position"):
                story.append(Paragraph(f"<i>primary:</i> {esc(dd.get('primary_position'))}", BODY))
            if dd.get("consultant_position"):
                story.append(Paragraph(f"<i>consultant:</i> {esc(dd.get('consultant_position'))}", BODY))
    else:
        story.append(Paragraph("<i>(no final result was produced)</i>", BODY))

    # Transcript
    story.append(Paragraph("Full Transcript", SECTION))
    for m in (messages or []):
        story.append(Paragraph(
            f"{esc(m.get('agent_name'))} &nbsp;·&nbsp; {esc(m.get('role'))} &nbsp;·&nbsp; {esc(m.get('message_type'))}",
            TURN))
        for label, value in _structured_fields(m):
            story.append(Paragraph(_prettify(label), LABEL))
            story.append(Paragraph(esc(value), BODY))
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="40%", thickness=0.4, color=HexColor("#e0e0e0")))

    story.append(Spacer(1, 18))
    story.append(HRFlowable(width="100%", thickness=0.7, color=HexColor("#cccccc")))
    story.append(Spacer(1, 4))
    story.append(Paragraph(f"Exported {esc(_now_iso())}. {_COPYRIGHT_LINE}", FOOTER))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                            topMargin=0.8 * inch, bottomMargin=0.8 * inch,
                            title="AI Conclave Deliberation Record", author="AI Switchboard")
    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

_MUTED = RGBColor(0x66, 0x66, 0x66)
_LABEL_GREY = RGBColor(0x88, 0x88, 0x88)
_HEAD_BLUE = RGBColor(0x2C, 0x57, 0x77)


def render_docx(task: dict, messages: list[dict], final_result: dict | None,
                agent_runs: list[dict]) -> bytes:
    doc = Document()

    # Title
    title = doc.add_heading("AI Conclave — Deliberation Record", level=0)
    sub = doc.add_paragraph()
    run = sub.add_run(task.get("user_request") or "")
    run.italic = True
    run.font.color.rgb = _MUTED

    # Metadata
    meta = doc.add_paragraph()
    meta_bits = [
        ("Task ID", task.get("id")),
        ("Mode", task.get("mode")),
        ("Status", task.get("status")),
        ("Participants", _agents_summary(task)),
        ("Invoked by", task.get("source_agent")),
        ("Created", task.get("created_at")),
        ("Updated", task.get("updated_at")),
    ]
    for i, (k, v) in enumerate(meta_bits):
        if v in (None, ""):
            continue
        r = meta.add_run(("" if i == 0 else "    ") + f"{k}: ")
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = _MUTED
        r2 = meta.add_run(str(v))
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = _MUTED

    # Decision (if any)
    decision = task.get("user_decision")
    if isinstance(decision, str) and decision.strip():
        doc.add_heading("Recorded Decision", level=1)
        if task.get("user_decided_at"):
            p = doc.add_paragraph()
            r = p.add_run(f"Recorded at {task['user_decided_at']}")
            r.italic = True
            r.font.size = Pt(8.5)
            r.font.color.rgb = _MUTED
        _add_multiline(doc, decision)

    # Final result
    doc.add_heading("Final Result", level=1)
    fr = final_result or {}
    if fr:
        p = doc.add_paragraph()
        r = p.add_run("Agreement level: ")
        r.bold = True
        p.add_run(str(fr.get("agreement_level") or ""))
        if fr.get("resolution_status"):
            p2 = doc.add_paragraph()
            r2 = p2.add_run("Resolution status: ")
            r2.bold = True
            p2.add_run(str(fr.get("resolution_status")))
        _add_multiline(doc, fr.get("final_answer") or "")
        for dd in (fr.get("disagreements") or []):
            doc.add_heading("Disagreements", level=2)
            p = doc.add_paragraph()
            p.add_run(str(dd.get("topic") or "")).bold = True
            if dd.get("primary_position"):
                pp = doc.add_paragraph()
                pp.add_run("primary: ").italic = True
                pp.add_run(str(dd.get("primary_position")))
            if dd.get("consultant_position"):
                pc = doc.add_paragraph()
                pc.add_run("consultant: ").italic = True
                pc.add_run(str(dd.get("consultant_position")))
    else:
        doc.add_paragraph("(no final result was produced)")

    # Transcript
    doc.add_heading("Full Transcript", level=1)
    for m in (messages or []):
        h = doc.add_heading(level=2)
        rh = h.add_run(f"{m.get('agent_name') or '?'}  ·  {m.get('role') or ''}  ·  {m.get('message_type') or ''}")
        rh.font.color.rgb = _HEAD_BLUE
        for label, value in _structured_fields(m):
            pl = doc.add_paragraph()
            rl = pl.add_run(_prettify(label))
            rl.bold = True
            rl.font.size = Pt(8)
            rl.font.color.rgb = _LABEL_GREY
            _add_multiline(doc, value)

    # Footer
    doc.add_paragraph()
    foot = doc.add_paragraph()
    rf = foot.add_run(f"Exported {_now_iso()}. {_COPYRIGHT_LINE}")
    rf.italic = True
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Page geometry
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_multiline(doc, text: str) -> None:
    """Add text to a DOCX, splitting on newlines so paragraphs survive."""
    text = "" if text is None else str(text)
    if not text.strip():
        doc.add_paragraph("")
        return
    for line in text.split("\n"):
        doc.add_paragraph(line)


__all__ = ["render_pdf", "render_docx", "filename_stem"]
